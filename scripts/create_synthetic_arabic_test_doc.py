from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


OUT = Path("fixtures/حكم_تجاري_2024_عينة_تجريبية.docx")


def add_paragraph(doc: Document, text: str, *, bold: bool = False) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(12)


def main() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)
    section.left_margin = Pt(72)
    section.right_margin = Pt(72)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("حكم تجاري تجريبي بشأن فسخ عقد توريد")
    title_run.bold = True
    title_run.font.name = "Arial"
    title_run.font.size = Pt(18)

    add_paragraph(doc, "تنبيه: هذا مستند اصطناعي بالكامل لغرض اختبار نظام إدارة المعرفة. لا يحتوي على بيانات عميل حقيقية أو أطراف حقيقية.", bold=True)
    add_paragraph(doc, "الجهة المصدرة: المحكمة التجارية بالرياض")
    add_paragraph(doc, "رقم القضية: 4521/1445")
    add_paragraph(doc, "التاريخ: 2024-03-12 (1445-09-02 هـ)")
    add_paragraph(doc, "الأطراف: شركة الفجر للتوريدات (مدعية) ضد مؤسسة النخبة للمقاولات (مدعى عليها).")
    add_paragraph(doc, "ملخص الوقائع: تعاقدت المدعية مع المدعى عليها على توريد معدات تجارية خلال مدة محددة. دفعت المدعية جزءاً من الثمن، إلا أن المدعى عليها تأخرت في التسليم وأخلت بالمواصفات الجوهرية المتفق عليها.")
    add_paragraph(doc, "دفعت المدعى عليها بأن التأخير كان بسبب ظروف تشغيلية عارضة، وطلبت مهلة إضافية لتنفيذ الالتزام. غير أن المحكمة لاحظت أن التأخير امتد لمدة طويلة وأن الإخلال مس جوهر العقد والغرض الاقتصادي منه.")
    add_paragraph(doc, "الأساس القانوني: استندت المحكمة إلى القواعد العامة في تنفيذ العقود بحسن نية، وإلى حق الدائن في طلب الفسخ عند الإخلال الجوهري بالالتزام التعاقدي متى تعذر تحقيق الغرض من العقد.")
    add_paragraph(doc, "منطوق الحكم: حكمت المحكمة بفسخ عقد التوريد وإلزام المدعى عليها برد المبالغ المستلمة، مع حفظ حق المدعية في المطالبة بالتعويض عن الضرر المثبت في دعوى مستقلة.")
    add_paragraph(doc, "كلمات مفتاحية: فسخ عقد، توريد، إخلال جوهري، تعويض، المحكمة التجارية، تنفيذ العقود.")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
